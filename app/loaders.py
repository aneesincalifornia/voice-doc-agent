import os
from pathlib import Path
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from typing import List

def load_document(path: str) -> List[Document]:
    """
    Load a document by file extension: .pdf, .docx, or .txt.

    Returns a list of Document objects with page_content and metadata.
    Raises FileNotFoundError if the file doesn't exist.
    Raises ValueError if the file type is unsupported.
    """
    path = Path(path)

    if not path.exists():
        raise FileNotFoundError(f"Document not found at: {path}")

    file_ext = path.suffix.lower()

    if file_ext == ".pdf":
        return load_pdf(str(path))
    elif file_ext == ".docx":
        return load_docx(str(path))
    elif file_ext == ".txt":
        return load_txt(str(path))
    elif file_ext == ".xlsx":
        return load_excel(str(path))
    else:
        raise ValueError(
            f"Unsupported file type: {file_ext}. "
            f"Supported types: .pdf, .docx, .txt, .xlsx"
        )

def load_pdf(pdf_path: str) -> List[Document]:
    """Load a PDF file using PyPDFLoader."""
    loader = PyPDFLoader(pdf_path)
    pages = loader.load()
    print(f"✓ Loaded {len(pages)} pages from PDF")
    return pages

def load_docx(docx_path: str) -> List[Document]:
    """Load a DOCX file."""
    from docx import Document as DocxDocument

    doc = DocxDocument(docx_path)

    # Extract all paragraphs into a single or per-paragraph Documents
    full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])

    if not full_text.strip():
        raise ValueError(f"No extractable text found in DOCX: {docx_path}")

    documents = [
        Document(
            page_content=full_text,
            metadata={
                "source": docx_path,
                "page": 0,
                "type": "docx"
            }
        )
    ]

    print(f"✓ Loaded DOCX document from {docx_path}")
    return documents

def load_txt(txt_path: str) -> List[Document]:
    """Load a plain text file."""
    with open(txt_path, "r", encoding="utf-8") as f:
        content = f.read()

    if not content.strip():
        raise ValueError(f"No text found in file: {txt_path}")

    documents = [
        Document(
            page_content=content,
            metadata={
                "source": txt_path,
                "page": 0,
                "type": "txt"
            }
        )
    ]

    print(f"✓ Loaded text file from {txt_path}")
    return documents

def load_excel(xlsx_path: str) -> List[Document]:
    """Load an Excel (.xlsx) file, converting each sheet to readable text."""
    from openpyxl import load_workbook

    workbook = load_workbook(xlsx_path, data_only=True)

    if not workbook.sheetnames:
        raise ValueError(f"No sheets found in Excel file: {xlsx_path}")

    documents = []

    for sheet_idx, sheet_name in enumerate(workbook.sheetnames):
        sheet = workbook[sheet_name]

        # Collect all rows as readable text
        rows_text = []
        has_data = False

        for row in sheet.iter_rows(min_row=1, values_only=True):
            # Skip completely empty rows
            if all(cell is None for cell in row):
                continue

            has_data = True

            # Format row as "Column_Name: value | Column_Name: value ..."
            # Use column letter as fallback if no header row is detected
            row_parts = []
            for col_idx, cell_value in enumerate(row, start=1):
                col_letter = chr(64 + col_idx) if col_idx <= 26 else f"Col{col_idx}"
                if cell_value is not None:
                    row_parts.append(f"{col_letter}: {cell_value}")

            if row_parts:
                rows_text.append(" | ".join(row_parts))

        if not has_data:
            continue  # Skip empty sheets

        sheet_content = "\n".join(rows_text)

        documents.append(
            Document(
                page_content=sheet_content,
                metadata={
                    "source": xlsx_path,
                    "page": sheet_idx,
                    "type": "xlsx",
                    "sheet_name": sheet_name
                }
            )
        )

    if not documents:
        raise ValueError(f"No data found in Excel file: {xlsx_path}")

    print(f"✓ Loaded {len(documents)} sheet(s) from Excel file")
    return documents
